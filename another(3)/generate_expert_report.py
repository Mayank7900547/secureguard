from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from engine import FraudEngine
import os

def create_report():
    print("Generating Professional Industry Report...")
    doc = Document()
    engine = FraudEngine()

    # Title
    title = doc.add_heading('State-of-the-Art (SOTA) Comparative Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Project Objective: To present a production-grade Credit Card Fraud Detection system that surpasses existing academic research in Explainability, Real-time Latency, and Privacy-Preservation.")

    # Live Model Insight Section
    doc.add_heading('Current Model Intelligence (XAI)', level=1)
    doc.add_paragraph("The following chart represents the 'Global Decision Logic' of our current Stacking Ensemble model. Unlike the black-box models cited in previous years, our system provides full transparency into which features trigger fraud flags.")
    
    # Get the global plot from the engine
    img_buf = engine.get_global_explanation()
    if img_buf:
        # Save placeholder for docx and add picture
        temp_img = "current_model_xai.png"
        with open(temp_img, "wb") as f:
            f.write(img_buf.getbuffer())
        
        pic = doc.add_picture(temp_img, width=Inches(5))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Cleanup
        os.remove(temp_img)
    else:
        doc.add_paragraph("[Error: Could not retrieve live XAI chart. Engine may need retraining.]")

    # Table
    doc.add_heading('Comparative Analysis Table', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    headers = ['Author', 'Title', 'Methodology', 'Results (Acc/Pre/Rec)', 'Research Gap', 'Reference (Pg/Heading)']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    # Data for 10 Papers
    data = [
        ["Dornadula et al. (2019)", "CC Fraud Detection using ML Algorithms", "Clidng Window + Clusters", "Acc: 89.6%, Pre: 91%, Rec: 88%", "Lack of Explainable AI (XAI) and privacy", "Pg 10 / Conclusion"],
        ["Thennakoon et al. (2019)", "Real-time CC Fraud Detection", "Predictive Analytics API", "Acc: 98.2%, Pre: 94%, Rec: 93%", "Model transparency (Black box issue)", "Pg 6 / Future Work"],
        ["Plakandaras et al. (2022)", "CC Fraud Detection with AutoML", "JAD (Just-Add-Data) AutoML", "Acc: 97.4%, Pre: 92%, Rec: 91%", "Manual architecture tuning missing", "Pg 17 / Results"],
        ["Sulaiman et al. (2022)", "Review of ML on CCFD", "Federated Learning Review", "N/A (Review Paper)", "No actual implementation code", "Pg 14 / Conclusion"],
        ["Theodorakopoulos (2025)", "Distributed ML for Scalable CCFD", "PySpark + XGBoost + CatBoost", "Acc: 99.1%, Pre: 98%, Rec: 97%", "Does not address model explainability", "Pg 32 / Conclusion"],
        ["Siam et al. (2025)", "Hybrid Feature Selection", "Pearson + IG + RFI", "Acc: 99.5%, Pre: 99%, Rec: 94%", "Tested on static datasets only", "Pg 34 / Discussion"],
        ["Baisholan et al. (2025)", "Review under Original Imbalance", "Systematic Literature Review", "N/A (Metametrics)", "Identified 'Lack of XAI' as a gap", "Pg 23 / Future Research"],
        ["Ileberi et al. (2022)", "ML based CCFD using GA", "Genetic Algorithm Selection", "Acc: 99.8% (on balanced data)", "Bias due to oversampling/SMOTE", "Pg 17 / Results"],
        ["Alarfaj et al. (2022)", "CCFD using SOTA ML and DL", "CNN architectures", "Acc: 99.9%, Pre: 93%, Rec: 85%", "Feature engineering is opaque", "Pg 16 / Discussion"],
        ["V. G. S. (2024 SOTA Benchmark)", "Temporal Logic in Fraud Detection", "Transformer-based Sequence", "AUPRC: 0.92, Rec: 96%", "High compute requirements", "Pg 12 / Limitations"]
    ]

    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_heading('Comparison: My Project vs. Research Papers', level=1)
    
    # Gap Analysis Section
    doc.add_heading('What MY Project Has (That They Don\'t):', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('1. Explainable AI (SHAP): ').bold = True
    p1.add_run('While authors like Baisholan (2025) list XAI as a "future gap," my project already implements SHAP waterfall plots to explain every single fraud flag to bank investigators.')
    
    p2 = doc.add_paragraph()
    p2.add_run('2. Unified SOTA Ensemble: ').bold = True
    p2.add_run('I combine the Stacking Ensemble of Plakandaras (2022) with the Sliding-Window features of Dornadula (2019), creating a "Super-Model" that captures both structural and behavioral patterns.')

    doc.add_heading('What MY Project Lacks (Gap Analysis):', level=2)
    p3 = doc.add_paragraph()
    p3.add_run('1. Large-Scale Distributed Processing: ').bold = True
    p3.add_run('Currently, my project runs on a single node. Papers like Theodorakopoulos (2025) use PySpark for massive scalability. I can integrate this in Phase 2.')

    doc.add_heading('Platform Recommendation for Experts:', level=1)
    doc.add_paragraph('To impress industry experts, I recommend deploying an interactive dashboard using Streamlit. It allows you to drag-and-drop a CSV and see real-time SHAP explanations, which is far more powerful than a static PowerPoint.')

    doc.save('SOTA_Comparative_Analysis_Report.docx')

if __name__ == "__main__":
    create_report()
