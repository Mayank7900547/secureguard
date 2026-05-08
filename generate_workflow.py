from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_workflow_doc():
    doc = Document()

    # --- Header Styling ---
    header = doc.add_heading('SOTA Fraud Detection: Project Workflow & Strategic Enhancements', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Author: Project Lead\nFocus: Industry-Expert Evaluation & SOTA Benchmark Alignment")

    # --- Section 1: Core Enhancements (Current) ---
    doc.add_heading('1. Current Project Enhancements', level=1)
    doc.add_paragraph("The project has transitioned from a standard classification model to a 'State-of-the-Art' (SOTA) expert system through the following enhancements:")

    enhancements = [
        ("Stacking Ensemble Architecture", "Instead of a single model, we use a meta-learner (Logistic Regression) that aggregates predictions from Random Forest and XGBoost. This minimizes bias and maximizes recall."),
        ("Explainable AI (SHAP) Layer", "We have bridged the 'Black Box Gap' by integrating SHAP. Every fraud flag is accompanied by a Waterfall Plot explaining the exact impact of each feature (Amount, Distance, etc.)."),
        ("Sliding-Window Feature Engineering", "The system calculates 'temporal' features (e.g., Average Spending over last 7 days) in real-time to detect behavioral shifts, not just static triggers.")
    ]

    for title, desc in enhancements:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # --- Section 2: Technical Workflow ---
    doc.add_heading('2. System Workflow', level=1)
    
    workflow_steps = [
        "Data Ingestion: Real-time acquisition via CSV upload or manual entry.",
        "Preprocessing: Normalization of transaction amounts and distance calculations.",
        "Ensemble Inference: Stacking models perform parallel prediction.",
        "XAI Generation: KernelExplainer extracts SHAP values for the specific transaction.",
        "Expert Dashboard: Results are visualized with high-confidence flags and trust-decoding charts."
    ]
    
    for step in workflow_steps:
        doc.add_paragraph(step, style='List Number')

    # --- Section 3: Planned Enhancements (Phase 2 Roadmap) ---
    doc.add_heading('3. Planned Future Enhancements (Phase 2)', level=1)
    doc.add_paragraph("To push the project beyond academic excellence into enterprise-grade readiness, the following enhancements are planned:")

    future = [
        ("Federated Learning Integration", "Allowing model training across decentralized bank servers without directly sharing sensitive PII (Personally Identifiable Information)."),
        ("Concept Drift Detection", "An automated monitoring layer that alerts developers when fraud patterns change (e.g., new types of e-commerce fraud) so the model can retrain."),
        ("Low-Latency API Gateway", "Transitioning the dashboard into a full FastAPI microservice for integration into mobile banking apps (Target Latency: <50ms).")
    ]

    for title, desc in future:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # Footer
    doc.add_paragraph("\n--- End of Workflow Documentation ---").alignment = WD_ALIGN_PARAGRAPH.CENTER

    fname = "Project_Workflow_and_Enhancements.docx"
    doc.save(fname)
    print(f"File saved successfully: {fname}")

if __name__ == "__main__":
    create_workflow_doc()
